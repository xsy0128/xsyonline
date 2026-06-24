import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = os.path.dirname(os.path.abspath(__file__))
doc = Document()

# ===== Page setup =====
for s in doc.sections:
    s.page_width = Cm(21); s.page_height = Cm(29.7)
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.18); s.right_margin = Cm(3.18)

# ===== Styles =====
style = doc.styles["Normal"]
style.font.name = "SimSun"
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

def set_font(run, name="SimSun", size=Pt(12), bold=False, color=None):
    run.font.name = name; run.font.size = size; run.bold = bold
    ea = "SimHei" if name == "SimHei" else name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), ea)
    if color: run.font.color.rgb = RGBColor(*color)

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    sz = {1: Pt(16), 2: Pt(14), 3: Pt(13)}.get(level, Pt(12))
    for r in h.runs: set_font(r, "SimHei", sz, True)
    return h

def P(doc, text, bold=False, indent=True, fs=Pt(12), align=None):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None: p.alignment = align
    run = p.add_run(text)
    set_font(run, "SimSun", fs, bold=bold)
    return p

def PL(doc, label):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"【请在此处粘贴{label}截图】")
    set_font(run, "SimSun", Pt(11), color=(150, 150, 150))
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("─" * 40)
    set_font(r2, "SimSun", Pt(8), color=(200, 200, 200))

def load_section(filename):
    path = os.path.join(BASE, "_report_sections", filename)
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    return None

def process_content(doc, text):
    lines = text.split("\n")
    for line in lines:
        line = line.strip()
        if not line: continue
        if line == "--pagebreak--": doc.add_page_break()
        elif line.startswith("[IMG:"):
            label = line[5:-1] if line.endswith("]") else line[5:]
            PL(doc, label)
        elif line.startswith("## "): H(doc, line[3:], level=1)
        elif line.startswith("### "): H(doc, line[4:], level=2)
        elif line.startswith("#### "): H(doc, line[5:], level=3)
        elif line.startswith("**"): P(doc, line[2:], bold=True, indent=True, fs=Pt(12))
        else: P(doc, line, indent=True, fs=Pt(12))

def add_page_numbers(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    pf = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = pf.add_run("第 "); set_font(r1, "SimSun", Pt(9))
    r2 = pf.add_run(); r2._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    r3 = pf.add_run(); r3._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    r4 = pf.add_run(); r4._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    r5 = pf.add_run(" 页"); set_font(r5, "SimSun", Pt(9))

def add_header(section, text):
    header = section.header
    header.is_linked_to_previous = False
    ph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = ph.add_run(text); set_font(r, "SimHei", Pt(8), color=(136, 136, 136))

# ===== TITLE PAGE =====
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("《Web前端开发》实验报告（一）")
set_font(run, "SimHei", Pt(22), True, color=None)
doc.add_paragraph()

# Info table
table = doc.add_table(rows=8, cols=2, style="Table Grid")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
info = [
    ("实验序号", "1"), ("实验项目名称", "个人主页制作 —— XSYonline个人网站"),
    ("学号", ""), ("姓名", "许诗怡"),
    ("专业、班", "信息管理与信息系统"), ("实验地点", "课外实验"),
    ("指导教师", "李胜"), ("时间", "2026.4.7"),
]
for row_idx, (label, value) in enumerate(info):
    cl = table.cell(row_idx, 0); cl.text = ""
    p = cl.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label); set_font(run, "SimSun", Pt(11), bold=True)
    cl._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF0F5" w:val="clear"/>'))
    cv = table.cell(row_idx, 1); cv.text = ""
    p = cv.paragraphs[0]; run = p.add_run(value); set_font(run, "SimSun", Pt(11))
for row in table.rows: row.cells[0].width = Cm(3.5); row.cells[1].width = Cm(12)
doc.add_paragraph()

# Add page numbers to cover
for s in doc.sections: add_page_numbers(s)
print("Phase 1: Title + Info table - done")

# ===== TOC in a new section =====
doc.add_section()
sec_toc = doc.sections[-1]
sec_toc.page_width = Cm(21); sec_toc.page_height = Cm(29.7)
sec_toc.top_margin = Cm(2.54); sec_toc.bottom_margin = Cm(2.54)
sec_toc.left_margin = Cm(3.18); sec_toc.right_margin = Cm(3.18)
add_header(sec_toc, "《Web前端开发》实验报告（一）")
add_page_numbers(sec_toc)

H(doc, "目  录", level=1)
p = doc.add_paragraph()
run = p.add_run("（用 Microsoft Word 打开后，右键此处 → 更新域 → 更新整个目录，即可自动生成目录和页码）")
set_font(run, "SimSun", Pt(10), color=(150, 150, 150))
p = doc.add_paragraph()
run = p.add_run(); run._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
r2 = p.add_run(); r2._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'))
r3 = p.add_run(); r3._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
r4 = p.add_run("（请在此处更新目录域）"); set_font(r4, "SimSun", Pt(11), color=(180, 180, 180))
r5 = p.add_run(); r5._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
print("Phase 2: TOC - done")

# ===== CONTENT SECTIONS =====
# Each content section gets its own docx section with header + page numbers
section_files = sorted([f for f in os.listdir(os.path.join(BASE, "_report_sections")) if f.endswith('.txt')])

for fname in section_files:
    print(f"Processing: {fname}")
    text = load_section(fname)
    if not text: continue

    # New docx section for each file
    doc.add_section()
    curr_sec = doc.sections[-1]
    curr_sec.page_width = Cm(21); curr_sec.page_height = Cm(29.7)
    curr_sec.top_margin = Cm(2.54); curr_sec.bottom_margin = Cm(2.54)
    curr_sec.left_margin = Cm(3.18); curr_sec.right_margin = Cm(3.18)
    add_header(curr_sec, "《Web前端开发》实验报告（一）  |  许诗怡  |  XSYonline")
    add_page_numbers(curr_sec)

    process_content(doc, text)

# ===== SAVE =====
output = os.path.join(BASE, "实验报告_许诗怡_XSYonline.docx")
doc.save(output)
print(f"Saved to: {output}")
print("Done!")
